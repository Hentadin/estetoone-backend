#!/usr/bin/env python3
"""Generate HEN-25 formal API documentation (DOCX + PDF)."""

from __future__ import annotations

import subprocess
import sys
from pathlib import Path

try:
    import yaml
    from docx import Document
    from docx.shared import Inches, Pt
except ImportError:
    print("Installing python-docx and pyyaml...", file=sys.stderr)
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "pyyaml", "-q"])
    import yaml
    from docx import Document
    from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
OPENAPI_PATH = ROOT / "docs" / "openapi.yaml"
OUTPUT_DIR = ROOT / "docs" / "formal"


def load_operations() -> list[dict]:
    spec = yaml.safe_load(OPENAPI_PATH.read_text(encoding="utf-8"))
    rows: list[dict] = []
    for path, methods in spec.get("paths", {}).items():
        for method, operation in methods.items():
            if method.startswith("x-"):
                continue
            tags = operation.get("tags") or ["other"]
            rows.append(
                {
                    "domain": tags[0],
                    "method": method.upper(),
                    "path": path,
                    "summary": operation.get("summary", ""),
                    "auth": infer_auth(operation),
                }
            )
    rows.sort(key=lambda r: (r["domain"], r["path"], r["method"]))
    return rows


def infer_auth(operation: dict) -> str:
    security = operation.get("security") or []
    if not security:
        return "Public"
    schemes = set()
    for requirement in security:
        schemes.update(requirement.keys())
    if "access-token" in schemes:
        return "Bearer JWT"
    if "refresh_token" in schemes:
        return "Cookie (refresh_token)"
    if "stripe-signature" in schemes:
        return "Stripe signature"
    return ", ".join(sorted(schemes))


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def build_document(operations: list[dict]) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    add_heading(doc, "EstetoOne API Documentation", 0)
    add_paragraph(doc, "HEN-25 — OpenAPI / Swagger deliverable")
    add_paragraph(doc, "Generated from estetoone-backend OpenAPI 3.x specification")

    add_heading(doc, "1. Environment setup", 1)
    add_paragraph(doc, "Backend (estetoone-backend):")
    for line in [
        "cp .env.example .env",
        "npm install && npx prisma migrate deploy && npm run prisma:seed",
        "npm run start:dev  →  http://localhost:3000",
    ]:
        doc.add_paragraph(line, style="List Bullet")
    add_paragraph(doc, "Frontend (stetho-ai-connect-25039, optional):")
    for line in [
        "npm install && npm run dev  →  http://localhost:8080",
        "Dev hub: http://localhost:8080/dev/api-testing",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    add_heading(doc, "2. Swagger UI access", 1)
    add_paragraph(doc, "Interactive docs with Try-it-out:")
    add_paragraph(doc, "URL: http://localhost:3000/api/docs", bold=True)
    add_paragraph(doc, "OpenAPI JSON: http://localhost:3000/api/docs-json")
    add_paragraph(doc, "Workflow:")
    for step in [
        "Open Swagger UI in browser",
        "POST /v1/auth/login with seed credentials (see section 5)",
        "Copy accessToken → Authorize → Bearer token",
        "Use Try-it-out on any endpoint",
    ]:
        doc.add_paragraph(step, style="List Number")

    add_heading(doc, "3. Endpoint inventory (51 operations)", 1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["Domain", "Method", "Path", "Auth", "Summary"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for op in operations:
        row = table.add_row().cells
        row[0].text = op["domain"]
        row[1].text = op["method"]
        row[2].text = op["path"]
        row[3].text = op["auth"]
        row[4].text = op["summary"]

    add_heading(doc, "4. Examples by domain", 1)
    examples = {
        "auth": (
            "POST /v1/auth/login",
            '{"email":"maria.silva@example.com","password":"password123"}',
            '{"accessToken":"...","user":{"id":"...","email":"...","role":"patient"}}',
        ),
        "health": ("GET /v1/health/live", "(no body)", '{"status":"ok","timestamp":"..."}'),
        "plans": ("GET /v1/plans", "(no body)", '[{"id":"basic","name":"Basic",...}]'),
        "consultations": (
            "POST /v1/consultations/schedule",
            '{"doctorId":"<uuid>","scheduledAt":"2026-12-01T10:00:00.000Z"}',
            '{"id":"...","status":"waiting","type":"SCHEDULED",...}',
        ),
        "payments": (
            "POST /v1/payments/setup-intent",
            "(Bearer required, no body)",
            '{"clientSecret":"seti_mock_...","setupIntentId":"seti_mock_..."}',
        ),
    }
    for domain, (endpoint, request, response) in examples.items():
        add_heading(doc, domain, 2)
        add_paragraph(doc, f"Endpoint: {endpoint}")
        add_paragraph(doc, f"Request: {request}")
        add_paragraph(doc, f"Response: {response}")

    add_heading(doc, "5. Seed test credentials", 1)
    doc.add_paragraph("Patient: maria.silva@example.com / password123", style="List Bullet")
    doc.add_paragraph("Doctor: dr.carlos@example.com / password123", style="List Bullet")

    add_heading(doc, "6. Mocked integrations & known gaps", 1)
    for item in [
        "Stripe: MockStripeClient when STRIPE_SECRET_KEY is empty",
        "Daily.co video: returns https://mock.daily.co/room-{consultationId}",
        "S3 documents: stored inline in database (s3Key column)",
        "CORS: not enabled in backend — use Swagger on :3000 or configure proxy",
        "Regenerate spec: npm run openapi:export in estetoone-backend",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    return doc


def main() -> None:
    if not OPENAPI_PATH.exists():
        raise SystemExit(f"Missing {OPENAPI_PATH}. Run npm run openapi:export first.")

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    operations = load_operations()
    doc = build_document(operations)
    docx_path = OUTPUT_DIR / "HEN-25-EstetoOne-API-Documentation.docx"
    doc.save(docx_path)
    print(f"Wrote {docx_path}")

    pdf_path = OUTPUT_DIR / "HEN-25-EstetoOne-API-Documentation.pdf"
    try:
        subprocess.run(
            [
                "libreoffice",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(OUTPUT_DIR),
                str(docx_path),
            ],
            check=True,
            capture_output=True,
        )
        print(f"Wrote {pdf_path}")
    except (FileNotFoundError, subprocess.CalledProcessError) as error:
        print(f"PDF conversion skipped: {error}", file=sys.stderr)


if __name__ == "__main__":
    main()
